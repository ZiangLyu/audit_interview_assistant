# -*- coding: utf-8 -*-
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import json
import requests
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import re
from typing import List, Optional
import io
import logging

app = FastAPI()

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("audit_interview")


# API配置
API_URL = "http://aihub.cofco.com:2885/v1/chat/completions"
API_KEY = "sk-HKx0az3Rg05DpULQ6b7951Bd7d5f4e28B17f7fB67fDf94Ab"

# 请求模型
class InterviewRequest(BaseModel):
    audit_type: str
    company_name: str
    interviewee_position: str
    interviewee_name: str
    transcript: str
    custom_prompt: Optional[str] = None

# 文件上传响应模型
class FileUploadResponse(BaseModel):
    content: str
    filename: str
    message: str

# 响应模型
class InterviewResponse(BaseModel):
    meeting_minutes: str
    audit_issues: List[str]
    required_materials: List[str]
    raw_response: str

def read_word_document(file_content: bytes) -> str:
    """读取Word文档内容"""
    try:
        # 将文件内容转换为BytesIO对象
        doc_stream = io.BytesIO(file_content)
        
        # 使用python-docx读取文档
        doc = Document(doc_stream)
        
        # 提取所有段落文本
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 只添加非空段落
                text_content.append(paragraph.text.strip())
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        # 合并所有文本内容
        full_content = "\n".join(text_content)
        
        if not full_content.strip():
            raise ValueError("文档内容为空")
            
        return full_content
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"读取Word文档失败: {str(e)}")

def call_llm(prompt: str) -> str:
    """调用大语言模型API"""
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": "Qwen2.5-72B-Instruct",
        "messages": [{
            "role": "user",
            "content": prompt
        }]
    }
    
    try:
        response = requests.post(API_URL, headers=headers, data=json.dumps(data), timeout=60)
        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content'].strip()
            return content
        else:
            raise HTTPException(status_code=response.status_code, detail=f"API请求失败: {response.text}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"API请求异常: {str(e)}")

def parse_llm_response(response: str) -> dict:
    """解析LLM响应，提取会议纪要、审计疑点和资料清单"""
    # 使用正则表达式分割不同部分
    sections = {
        'meeting_minutes': '',
        'audit_issues': [],
        'required_materials': []
    }
    
    # 查找会议纪要部分
    minutes_pattern = r'(?:会议纪要|访谈记录|会议记录)(.*?)(?=审计疑点|审计问题|$)'
    minutes_match = re.search(minutes_pattern, response, re.DOTALL | re.IGNORECASE)
    if minutes_match:
        sections['meeting_minutes'] = minutes_match.group(1).strip()
    else:
        # 如果没有明确的会议纪要标题，取第一部分作为会议纪要
        first_section_end = re.search(r'(?=审计疑点|审计问题|审计资料)', response)
        if first_section_end:
            sections['meeting_minutes'] = response[:first_section_end.start()].strip()
    
    # 查找审计疑点部分
    issues_pattern = r'(?:审计疑点|审计问题)(.*?)(?=审计资料|所需资料|资料清单|$)'
    issues_match = re.search(issues_pattern, response, re.DOTALL | re.IGNORECASE)
    if issues_match:
        issues_text = issues_match.group(1).strip()
        # 提取列表项
        issues_list = re.findall(r'[-•\d]+\.?\s*(.+?)(?=\n[-•\d]+\.?\s*|\n\n|$)', issues_text, re.DOTALL)
        sections['audit_issues'] = [issue.strip() for issue in issues_list if issue.strip()]
    
    # 查找资料清单部分
    materials_pattern = r'(?:审计资料|所需资料|资料清单)(.*?)$'
    materials_match = re.search(materials_pattern, response, re.DOTALL | re.IGNORECASE)
    if materials_match:
        materials_text = materials_match.group(1).strip()
        # 提取列表项
        materials_list = re.findall(r'[-•\d]+\.?\s*(.+?)(?=\n[-•\d]+\.?\s*|\n\n|$)', materials_text, re.DOTALL)
        sections['required_materials'] = [material.strip() for material in materials_list if material.strip()]
    
    return sections

def create_word_document(data: dict, filename: str) -> str:
    """创建Word文档"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('审计访谈记录', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加基本信息
    doc.add_paragraph(f'审计类型：{data.get("audit_type", "")}')
    doc.add_paragraph(f'被审计单位：{data.get("company_name", "")}')
    doc.add_paragraph(f'被访谈人：{data.get("interviewee_name", "")} ({data.get("interviewee_position", "")})')
    doc.add_paragraph(f'日期：{datetime.now().strftime("%Y年%m月%d日")}')
    doc.add_paragraph()
    
    # 添加会议纪要
    doc.add_heading('一、访谈记录', level=1)
    # 处理Markdown格式的会议纪要
    minutes_html = markdown.markdown(data.get("meeting_minutes", ""))
    # 简单处理，去除HTML标签
    minutes_text = re.sub('<.*?>', '', minutes_html)
    doc.add_paragraph(minutes_text)
    
    # 添加审计疑点
    doc.add_heading('二、审计疑点', level=1)
    for i, issue in enumerate(data.get("audit_issues", []), 1):
        doc.add_paragraph(f'{i}. {issue}')
    
    # 添加资料清单
    doc.add_heading('三、所需审计资料清单', level=1)
    for i, material in enumerate(data.get("required_materials", []), 1):
        doc.add_paragraph(f'{i}. {material}')
    
    # 保存文档
    filepath = f"output/{filename}"
    os.makedirs("output", exist_ok=True)
    doc.save(filepath)
    
    return filepath

@app.post("/api/upload-word", response_model=FileUploadResponse)
async def upload_word_document(file: UploadFile = File(...)):
    """上传并读取Word文档"""
    # 检查文件类型
    if not file.filename.lower().endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="只支持.docx和.doc格式的Word文档")
    
    try:
        # 读取文件内容
        file_content = await file.read()
        
        # 读取Word文档内容
        content = read_word_document(file_content)
        
        return FileUploadResponse(
            content=content,
            filename=file.filename,
            message="Word文档读取成功"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件处理失败: {str(e)}")

@app.post("/api/audit_interview/process-interview", response_model=InterviewResponse)
async def process_interview(request: InterviewRequest):
    """处理访谈记录"""

    logger.info('process_interview')
    # 构建提示词
    if request.custom_prompt:
        prompt = request.custom_prompt
    else:
        prompt = f"""假如你是一名资深的审计人员，你在开展{request.audit_type}审计，被审计单位是{request.company_name}，被访谈人：{request.company_name}{request.interviewee_position}{request.interviewee_name}员工。

请你根据以下访谈内容进行整理访谈会议纪要，并以Markdown形式输出整理后的访谈记录、审计疑点以及审计资料清单。

请按照以下格式输出：

## 会议纪要
[整理后的访谈内容]

## 审计疑点
1. [疑点1]
2. [疑点2]
...

## 审计资料清单
1. [资料1]
2. [资料2]
...

访谈内容：
{request.transcript}"""
    
    # 调用LLM
    llm_response = call_llm(prompt)
    logger.info('call llm')
    
    # 解析响应
    parsed = parse_llm_response(llm_response)
    
    return InterviewResponse(
        meeting_minutes=parsed['meeting_minutes'],
        audit_issues=parsed['audit_issues'],
        required_materials=parsed['required_materials'],
        raw_response=llm_response
    )

@app.post("/api/export-word")
async def export_word(request: dict):
    """导出Word文档"""
    filename = f"audit_interview_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = create_word_document(request, filename)
    
    return FileResponse(
        path=filepath,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.get("/")
async def root():
    return {"message": "审计访谈助手API"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)